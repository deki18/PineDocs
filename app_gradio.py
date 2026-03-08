import os
import base64
import io
import hashlib
from typing import List, Tuple, Optional
from datetime import datetime
from pathlib import Path

import gradio as gr
from dotenv import load_dotenv
import requests
import fitz  # PyMuPDF

from pinecone import Pinecone, ServerlessSpec, CloudProvider, AwsRegion
from openai import OpenAI

load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
EMBEDDING_API_KEY = os.getenv("EMBEDDING_API_KEY") or os.getenv("VECTORENGINE_API_KEY")
EMBEDDING_BASE_URL = os.getenv("EMBEDDING_BASE_URL", "https://api.vectorengine.ai/v1")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-large")
EMBEDDING_DIMENSION = int(os.getenv("EMBEDDING_DIMENSION", "3072"))

DOCS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
os.makedirs(DOCS_DIR, exist_ok=True)


# ============ 核心功能函数 ============

def generate_safe_id(filename: str, batch_id: int, chunk_id: int) -> str:
    """生成 ASCII 安全的 Vector ID"""
    ext = Path(filename).suffix
    name_hash = hashlib.md5(filename.encode('utf-8')).hexdigest()[:8]
    safe_id = f"doc_{name_hash}_b{batch_id}_c{chunk_id}{ext}"
    return safe_id


def get_embed_client(api_key=None, base_url=None):
    """获取嵌入客户端"""
    key = api_key or EMBEDDING_API_KEY
    url = base_url or EMBEDDING_BASE_URL
    
    if not key:
        return None
    return OpenAI(
        base_url=url,
        api_key=key,
    )


def get_pinecone_index(index_name: str, dimension=None):
    """获取或创建 Pinecone 索引"""
    if not PINECONE_API_KEY:
        return None, "请在.env文件中设置PINECONE_API_KEY"

    try:
        pc = Pinecone(api_key=PINECONE_API_KEY)
        existing_indexes = [idx["name"] for idx in pc.list_indexes()]
        
        dim = dimension or EMBEDDING_DIMENSION
        
        if index_name not in existing_indexes:
            pc.create_index(
                name=index_name,
                dimension=dim,
                metric="cosine",
                spec=ServerlessSpec(
                    cloud=CloudProvider.AWS,
                    region=AwsRegion.US_EAST_1
                ),
            )
        
        return pc.Index(index_name), None
    except Exception as e:
        return None, f"获取索引失败: {str(e)}"


def embed_texts(texts: List[str], model=None, api_key=None, base_url=None) -> Tuple[List[List[float]], str]:
    """将文本转换为向量"""
    client = get_embed_client(api_key=api_key, base_url=base_url)
    if not client:
        return [], "未配置嵌入模型API密钥"

    try:
        model_name = model or EMBEDDING_MODEL
        response = client.embeddings.create(
            model=model_name,
            input=texts,
        )
        return [item.embedding for item in response.data], None
    except Exception as e:
        return [], f"嵌入失败: {str(e)}"


def split_text(text: str, max_chars: int = 1000, overlap: int = 200) -> List[str]:
    """将文本分割成块"""
    if max_chars <= overlap:
        raise ValueError("max_chars must be greater than overlap")

    chunks: List[str] = []
    start = 0
    length = len(text)

    while start < length:
        end = min(start + max_chars, length)
        
        if end < length:
            last_newline = text.rfind('\n', start, end)
            last_period = text.rfind('。', start, end)
            last_exclamation = text.rfind('！', start, end)
            last_question = text.rfind('？', start, end)
            last_semicolon = text.rfind('；', start, end)
            
            best_break = max(last_newline, last_period, last_exclamation, last_question, last_semicolon)
            
            if best_break > start + max_chars // 2:
                end = best_break + 1
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        if end == length:
            break
        
        start = max(end - overlap, 0)

    return chunks


def split_markdown_by_headers(markdown_text: str) -> List[Tuple[str, str]]:
    """
    按Markdown标题层级切分文本
    返回: [(header, content), ...]
    """
    import re
    
    # 匹配 # 到 ###### 各级标题 (1-6级标题)
    header_pattern = r'^(#{1,6})\s+(.+)$'
    
    sections = []
    current_header = "文档开头"
    current_content = []
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        match = re.match(header_pattern, line.strip())
        if match:
            # 保存上一个章节
            if current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    sections.append((current_header, content))
            # 开始新章节
            current_header = match.group(2).strip()
            current_content = []
        else:
            current_content.append(line)
    
    # 保存最后一个章节
    if current_content:
        content = '\n'.join(current_content).strip()
        if content:
            sections.append((current_header, content))
    
    # 如果没有找到任何标题，将整个文本作为一个章节
    if not sections and markdown_text.strip():
        sections.append(("全文", markdown_text.strip()))
    
    return sections


def is_table_line(line: str) -> bool:
    """检查是否是Markdown表格行"""
    stripped = line.strip()
    # 表格行特征: 以 | 开头和结尾，或者包含 | 分隔符
    return stripped.startswith('|') and stripped.endswith('|') and '|' in stripped[1:-1]


def is_table_separator(line: str) -> bool:
    """检查是否是表格分隔符行 (如 |---|---|)"""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    # 去除首尾的 |，检查中间是否只包含 -、: 和 |
    inner = stripped[1:-1]
    return all(c in '-:| ' for c in inner)


def extract_tables(text: str) -> List[Tuple[str, str]]:
    """
    从文本中提取表格
    返回: [(table_header, table_content), ...]
    """
    lines = text.split('\n')
    tables = []
    current_table = []
    table_header = ""
    in_table = False
    
    for i, line in enumerate(lines):
        if is_table_line(line) and not is_table_separator(line):
            if not in_table:
                # 表格开始，尝试获取表头
                in_table = True
                current_table = [line]
                # 向上查找可能的表格标题（非空行）
                for j in range(i-1, max(i-5, -1), -1):
                    prev_line = lines[j].strip()
                    if prev_line and not is_table_line(prev_line):
                        table_header = prev_line
                        break
            else:
                current_table.append(line)
        elif in_table:
            if line.strip() == '' or is_table_separator(line):
                current_table.append(line)
            else:
                # 表格结束
                if len(current_table) >= 2:  # 至少表头+一行数据
                    table_content = '\n'.join(current_table).strip()
                    if table_content:
                        tables.append((table_header, table_content))
                in_table = False
                current_table = []
                table_header = ""
    
    # 处理文档末尾的表格
    if in_table and len(current_table) >= 2:
        table_content = '\n'.join(current_table).strip()
        if table_content:
            tables.append((table_header, table_content))
    
    return tables


def split_pdf_markdown(markdown_text: str, chunk_size: int = 512, chunk_overlap: int = 100) -> List[Tuple[str, str, str]]:
    """
    PDF专用三层级切分逻辑
    
    Level 1: 按 ## 二级标题切分
    Level 2: 如果内容超过 chunk_size，按段落(\n\n)切分
    Level 3: 如果段落还超长，按字符强制切分
    
    返回: [(header, content, content_type), ...]
        content_type: 'text' 或 'table'
    """
    import re
    
    results = []
    
    # Level 1: 按标题粗切
    sections = split_markdown_by_headers(markdown_text)
    print(f"[PDF Split] Found {len(sections)} sections by headers")
    
    for header, section_content in sections:
        # 先提取表格，单独处理
        tables = extract_tables(section_content)
        table_ranges = []
        
        # 标记表格位置，避免重复切分
        for table_header_text, table_content in tables:
            # 表格作为一个整体
            if len(table_content) <= chunk_size * 2:  # 不太长的表格直接保留
                results.append((f"{header} - {table_header_text}" if table_header_text else header, 
                              table_content, 'table'))
            else:
                # 长表格按行切分，每行都带表头
                table_lines = table_content.split('\n')
                if len(table_lines) >= 2:
                    # 第一行是表头
                    header_line = table_lines[0]
                    separator = table_lines[1] if is_table_separator(table_lines[1]) else ""
                    data_rows = table_lines[2:] if separator else table_lines[1:]
                    
                    # 每 chunk_size/2 字符左右切分一次
                    current_chunk = [header_line]
                    if separator:
                        current_chunk.append(separator)
                    current_size = len(header_line) + len(separator)
                    
                    for row in data_rows:
                        if current_size + len(row) > chunk_size and current_chunk:
                            # 保存当前块
                            results.append((f"{header} - {table_header_text} (表格片段)" if table_header_text else f"{header} (表格片段)",
                                          '\n'.join(current_chunk), 'table'))
                            # 新开一块，带上表头
                            current_chunk = [header_line]
                            if separator:
                                current_chunk.append(separator)
                            current_size = len(header_line) + len(separator)
                        
                        current_chunk.append(row)
                        current_size += len(row) + 1
                    
                    # 保存最后一块
                    if len(current_chunk) > (2 if separator else 1):
                        results.append((f"{header} - {table_header_text} (表格片段)" if table_header_text else f"{header} (表格片段)",
                                      '\n'.join(current_chunk), 'table'))
        
        # 移除表格内容，处理剩余文本
        text_without_tables = section_content
        for _, table_content in tables:
            text_without_tables = text_without_tables.replace(table_content, '')
        text_without_tables = re.sub(r'\n{3,}', '\n\n', text_without_tables).strip()
        
        if not text_without_tables:
            continue
        
        # Level 2 & 3: 对非表格内容进行细切
        if len(text_without_tables) <= chunk_size:
            # 内容较短，直接保留
            results.append((header, text_without_tables, 'text'))
        else:
            # 按段落切分
            paragraphs = text_without_tables.split('\n\n')
            current_chunk = []
            current_size = 0
            chunk_index = 1
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                para_size = len(para)
                
                # 如果单个段落就超过 chunk_size，需要强制切分
                if para_size > chunk_size:
                    # 先保存当前累积的内容
                    if current_chunk:
                        content = '\n\n'.join(current_chunk)
                        results.append((f"{header} (片段{chunk_index})", content, 'text'))
                        chunk_index += 1
                        current_chunk = []
                        current_size = 0
                    
                    # 强制切分这个长段落
                    start = 0
                    while start < para_size:
                        end = min(start + chunk_size, para_size)
                        # 尝试在句子边界切分
                        if end < para_size:
                            for sep in ['。', '！', '？', '；', '\n']:
                                pos = para.rfind(sep, start, end)
                                if pos > start + chunk_size // 2:
                                    end = pos + 1
                                    break
                        
                        sub_para = para[start:end].strip()
                        if sub_para:
                            results.append((f"{header} (片段{chunk_index})", sub_para, 'text'))
                            chunk_index += 1
                        
                        start = end - chunk_overlap if end < para_size else end
                else:
                    # 检查加入当前块是否会超限
                    if current_size + para_size + 2 > chunk_size and current_chunk:
                        # 保存当前块
                        content = '\n\n'.join(current_chunk)
                        results.append((f"{header} (片段{chunk_index})", content, 'text'))
                        chunk_index += 1
                        
                        # 新块，带上重叠内容（上一段的最后一部分）
                        if current_chunk and chunk_overlap > 0:
                            overlap_text = current_chunk[-1][-chunk_overlap:] if len(current_chunk[-1]) > chunk_overlap else current_chunk[-1]
                            current_chunk = [overlap_text, para]
                            current_size = len(overlap_text) + para_size + 2
                        else:
                            current_chunk = [para]
                            current_size = para_size
                    else:
                        current_chunk.append(para)
                        current_size += para_size + 2
            
            # 保存最后一个块
            if current_chunk:
                content = '\n\n'.join(current_chunk)
                if chunk_index == 1:
                    results.append((header, content, 'text'))
                else:
                    results.append((f"{header} (片段{chunk_index})", content, 'text'))
    
    print(f"[PDF Split] Generated {len(results)} chunks")
    return results


def extract_text_from_file(file_obj, ollama_url: str = "http://localhost:11434", ocr_model: str = "my-glm-ocr:latest") -> Tuple[str, str]:
    """从文件中提取文本"""
    try:
        # Gradio 的 File 组件返回的是文件路径（字符串）或文件对象
        if isinstance(file_obj, str):
            # 如果是字符串路径
            filename = file_obj
            with open(file_obj, 'rb') as f:
                content = f.read()
        elif hasattr(file_obj, 'name'):
            # 如果是文件对象
            filename = file_obj.name
            if hasattr(file_obj, 'read'):
                content = file_obj.read()
            else:
                # 可能是路径对象
                with open(file_obj.name, 'rb') as f:
                    content = f.read()
        else:
            return "", f"无效的文件对象类型: {type(file_obj)}"
        
        ext = Path(filename).suffix.lower()
        
        # 图片文件使用OCR
        if ext in ['.jpg', '.jpeg', '.png', '.bmp']:
            return ocr_image(content, ollama_url, ocr_model)
        
        # PDF文件
        elif ext == '.pdf':
            return pdf_to_text(content, ollama_url, ocr_model, filename)
        
        # Word文档
        elif ext in ['.doc', '.docx']:
            return extract_word_text(content, ext)
        
        # Excel文件
        elif ext == '.xlsx':
            return extract_excel_text(content)
        
        # 文本文件
        elif ext in ['.txt', '.md']:
            try:
                return content.decode('utf-8'), None
            except:
                try:
                    return content.decode('gbk'), None
                except:
                    return content.decode('latin-1'), None
        
        else:
            return "", f"不支持的文件格式: {ext}"
            
    except Exception as e:
        return "", f"提取文本失败: {str(e)}"


def ocr_image(image_bytes: bytes, ollama_url: str, model_name: str) -> Tuple[str, str]:
    """使用Ollama进行OCR识别"""
    try:
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        response = requests.post(
            f"{ollama_url}/api/generate",
            json={
                "model": model_name,
                "prompt": "请识别图片中的文字内容，直接输出识别到的文字，不要添加任何解释。",
                "images": [image_base64],
                "stream": False
            },
            timeout=120
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get("response", ""), None
        else:
            return "", f"OCR请求失败: {response.status_code}"
    except Exception as e:
        return "", f"OCR识别失败: {str(e)}"


def pdf_to_text(pdf_bytes: bytes, ollama_url: str, model_name: str, filename: str) -> Tuple[str, str]:
    """将PDF转换为文本，如果没有文本层则使用OCR"""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        all_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text.strip():
                all_text.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
        
        doc.close()
        
        # 如果没有提取到文本，可能是扫描版PDF，使用OCR
        if not all_text:
            print(f"PDF没有文本层，尝试使用OCR识别: {filename}")
            return pdf_to_text_with_ocr(pdf_bytes, ollama_url, model_name, filename)
        
        return "\n\n".join(all_text), None
    except Exception as e:
        return "", f"PDF处理失败: {str(e)}"


def pdf_to_text_with_ocr(pdf_bytes: bytes, ollama_url: str, model_name: str, filename: str = "") -> Tuple[str, str]:
    """使用OCR识别PDF中的文字（用于扫描版PDF）"""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        all_content = []
        
        print(f"开始OCR识别: {filename}，共 {total_pages} 页")
        
        # 保持高分辨率以获得更好的OCR效果
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(total_pages):
            print(f"  处理第 {page_num + 1}/{total_pages} 页...")
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # 对图片进行OCR识别
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            
            # 根据模型选择不同的调用方式
            if model_name == "my-PaddleOCR-VL:0.9b":
                response = requests.post(
                    f"{ollama_url}/api/chat",
                    json={
                        "model": "my-PaddleOCR-VL:0.9b",
                        "messages": [
                            {
                                "role": "user",
                                "content": "请识别这张图片中的文字内容。",
                                "images": [img_base64]
                            }
                        ],
                        "stream": False,
                        "options": {"temperature": 0.1, "num_predict": 4096}
                    },
                    timeout=300
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("message", {}).get("content", "")
                    if content:
                        all_content.append(f"--- 第 {page_num + 1} 页 ---\n{content}")
            else:
                response = requests.post(
                    f"{ollama_url}/api/generate",
                    json={
                        "model": "my-glm-ocr:latest",
                        "prompt": "Text Recognition:",
                        "images": [img_base64],
                        "stream": False
                    },
                    timeout=300
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("response", "")
                    if content:
                        all_content.append(f"--- 第 {page_num + 1} 页 ---\n{content}")
        
        doc.close()
        
        if all_content:
            return "\n\n".join(all_content), None
        else:
            return "", "OCR识别失败，未能提取到任何文字"
            
    except Exception as e:
        return "", f"OCR识别失败: {str(e)}"


def extract_word_text(file_bytes: bytes, ext: str) -> Tuple[str, str]:
    """从Word文档提取文本"""
    try:
        if ext == '.docx':
            # 处理 .docx 文件 (ZIP格式的XML)
            from docx import Document
            
            # 检查文件大小
            file_size = len(file_bytes)
            print(f"[DEBUG] DOCX file size: {file_size} bytes")
            
            if file_size == 0:
                return "", "文件大小为0，可能是空文件"
            
            doc = Document(io.BytesIO(file_bytes))
            
            all_text = []
            
            # 提取段落文本
            para_count = len(doc.paragraphs)
            print(f"[DEBUG] Number of paragraphs: {para_count}")
            
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text:
                    all_text.append(para.text)
                    if i < 5:  # 只打印前5个段落用于调试
                        print(f"[DEBUG] Paragraph {i}: {para_text[:100]}...")
            
            # 提取表格中的文本
            table_count = len(doc.tables)
            print(f"[DEBUG] Number of tables: {table_count}")
            
            for table_idx, table in enumerate(doc.tables):
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        all_text.append(" | ".join(row_texts))
            
            # 提取页眉中的文本
            section_count = len(doc.sections)
            print(f"[DEBUG] Number of sections: {section_count}")
            
            for section in doc.sections:
                try:
                    header = section.header
                    for para in header.paragraphs:
                        if para.text.strip():
                            all_text.append(para.text)
                except Exception as e:
                    print(f"[DEBUG] Header extraction error: {e}")
                
                # 提取页脚中的文本
                try:
                    footer = section.footer
                    for para in footer.paragraphs:
                        if para.text.strip():
                            all_text.append(para.text)
                except Exception as e:
                    print(f"[DEBUG] Footer extraction error: {e}")
            
            # 尝试从文档的XML中直接提取所有文本（备用方案）
            if not all_text:
                print("[DEBUG] Trying alternative text extraction from XML...")
                try:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import qn
                    
                    # 获取文档的XML元素
                    xml_content = doc.element.xml
                    # 简单的文本提取：找到所有<w:t>标签的内容
                    import re
                    text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                    alt_text = [t.strip() for t in text_matches if t.strip()]
                    if alt_text:
                        all_text.extend(alt_text)
                        print(f"[DEBUG] Alternative extraction found {len(alt_text)} text segments")
                except Exception as e:
                    print(f"[DEBUG] Alternative extraction failed: {e}")
            
            result = "\n".join(all_text)
            print(f"[DEBUG] Total extracted text length: {len(result)} characters")
            
            if not result.strip():
                return "", "文档内容为空（python-docx无法提取内容，可能是特殊格式的文档）"
            return result, None
        elif ext == '.doc':
            # 处理 .doc 文件 (旧版Word二进制格式)
            # 在Windows上使用Word COM接口
            try:
                import win32com.client
                import pythoncom
                import tempfile
                import os
                
                # 初始化COM
                pythoncom.CoInitialize()
                
                # 创建临时文件
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_doc:
                    temp_doc.write(file_bytes)
                    temp_doc_path = temp_doc.name
                
                temp_docx_path = temp_doc_path + 'x'
                
                try:
                    # 使用Word将.doc转换为.docx
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc_obj = word.Documents.Open(temp_doc_path)
                    doc_obj.SaveAs(temp_docx_path, FileFormat=16)  # 16 = wdFormatXMLDocument
                    doc_obj.Close()
                    word.Quit()
                    
                    # 读取转换后的.docx文件
                    from docx import Document
                    doc = Document(temp_docx_path)
                    
                    all_text = []
                    
                    # 提取段落文本
                    for para in doc.paragraphs:
                        if para.text.strip():
                            all_text.append(para.text)
                    
                    # 提取表格中的文本
                    for table in doc.tables:
                        for row in table.rows:
                            row_texts = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_texts.append(cell_text)
                            if row_texts:
                                all_text.append(" | ".join(row_texts))
                    
                    # 提取页眉页脚
                    for section in doc.sections:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                all_text.append(para.text)
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                all_text.append(para.text)
                    
                    result = "\n".join(all_text)
                    if not result.strip():
                        return "", "文档内容为空"
                    return result, None
                    
                finally:
                    # 清理临时文件
                    pythoncom.CoUninitialize()
                    if os.path.exists(temp_doc_path):
                        os.remove(temp_doc_path)
                    if os.path.exists(temp_docx_path):
                        os.remove(temp_docx_path)
                        
            except ImportError:
                return "", "无法处理 .doc 文件。请安装 pywin32: pip install pywin32"
            except Exception as e:
                return "", f".doc文件处理失败: {str(e)}"
        else:
            return "", f"不支持的Word格式: {ext}"
    except Exception as e:
        return "", f"Word文档处理失败: {str(e)}"


def extract_excel_text(file_bytes: bytes) -> Tuple[str, str]:
    """从Excel文件提取文本"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes))
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"--- 工作表: {sheet_name} ---"]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    sheet_text.append(row_text)
            
            all_text.append("\n".join(sheet_text))
        
        return "\n\n".join(all_text), None
    except Exception as e:
        return "", f"Excel处理失败: {str(e)}"


import sys

MAX_REQUEST_SIZE = 2 * 1024 * 1024

def estimate_vectors_size(vectors):
    size = 0
    for vec in vectors:
        size += sys.getsizeof(vec.get("id", ""))
        size += sys.getsizeof(vec.get("values", [])) * 8
        metadata = vec.get("metadata", {})
        for k, v in metadata.items():
            size += sys.getsizeof(k)
            size += sys.getsizeof(v)
    return int(size * 1.2)

def smart_upsert(index, vectors, namespace, results=None, depth=0):
    if not vectors:
        return 0
    
    estimated_size = estimate_vectors_size(vectors)
    
    if estimated_size <= MAX_REQUEST_SIZE or len(vectors) == 1:
        try:
            index.upsert(vectors=vectors, namespace=namespace)
            if results is not None and depth > 0:
                results.append(f"    ↳ 子批次上传 {len(vectors)} 个向量 ({estimated_size / 1024 / 1024:.2f}MB)")
            return len(vectors)
        except Exception as e:
            if "exceeds the maximum" in str(e) and len(vectors) > 1:
                pass
            else:
                raise e
    
    mid = len(vectors) // 2
    first_half = vectors[:mid]
    second_half = vectors[mid:]
    
    if results is not None:
        results.append(f"    ⚠️ 批次过大 ({estimated_size / 1024 / 1024:.2f}MB > 2MB)，自动分割为 {len(first_half)} + {len(second_half)} 个向量")
    
    count = 0
    count += smart_upsert(index, first_half, namespace, results, depth + 1)
    count += smart_upsert(index, second_half, namespace, results, depth + 1)
    
    return count

def batched(items, batch_size=10):
    for i in range(0, len(items), batch_size):
        yield items[i:i + batch_size]


def pdf_to_markdown_with_ollama(pdf_bytes: bytes, ollama_url: str = "http://localhost:11434", model_name: str = "my-glm-ocr:latest", filename: str = "document.pdf") -> Tuple[str, str]:
    """使用Ollama的OCR模型将PDF转换为Markdown
    
    支持的模型:
    - my-PaddleOCR-VL:0.9b: PaddleOCR-VL-1.5 (推荐，效果更好)
    - my-glm-ocr:latest: GLM-OCR模型
    
    Args:
        pdf_bytes: PDF文件的字节数据
        ollama_url: Ollama服务地址
        model_name: 使用的OCR模型名称
        filename: 原始文件名，用于保存Markdown文件
    
    Returns:
        (markdown_content, error_message)
    """
    try:
        # 使用PyMuPDF将PDF转换为图片列表
        try:
            # 打开PDF
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            images = []
            
            # 设置缩放比例 (zoom=2.0 相当于 144 DPI, zoom=3.0 相当于 216 DPI)
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                # 将页面渲染为图片
                pix = page.get_pixmap(matrix=mat)
                # 转换为PNG数据
                img_data = pix.tobytes("png")
                images.append(img_data)
            
            pdf_document.close()
            
        except Exception as e:
            return "", f"PDF转图片失败: {str(e)}"
        
        if not images:
            return "", "PDF中没有找到页面"
        
        all_content = []
        
        # 逐页处理
        for page_num, img_data in enumerate(images, start=1):
            # 将图片转换为base64
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            
            # 根据模型选择不同的调用方式
            if model_name == "my-PaddleOCR-VL:0.9b":
                # PaddleOCR-VL-1.5 使用chat API
                response = requests.post(
                    f"{ollama_url}/api/chat",
                    json={
                        "model": "my-PaddleOCR-VL:0.9b",
                        "messages": [
                            {
                                "role": "user",
                                "content": "请识别这张图片中的文字内容，并以Markdown格式输出。保留原有的段落结构和格式。",
                                "images": [img_base64]
                            }
                        ],
                        "stream": False,
                        "options": {
                            "temperature": 0.1,
                            "num_predict": 4096
                        }
                    },
                    timeout=300
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("message", {}).get("content", "")
                    if not content:
                        content = result.get("response", "")
                    if content:
                        all_content.append(f"## 第 {page_num} 页\n\n{content}")
            else:
                # GLM-OCR 使用Ollama原生 /api/generate 端点
                response = requests.post(
                    f"{ollama_url}/api/generate",
                    json={
                        "model": "my-glm-ocr:latest",
                        "prompt": """请分析图片内容并转换为 Markdown 格式。
要求：
1. 遇到表格，必须使用 Markdown 表格语法（| 列1 | 列2 |）
2. 遇到标题，使用 #, ##, ### 等标识层级
3. 保留原文的段落结构和列表格式
4. 不要包含"这是一张图片"等无关解释，直接输出 Markdown 内容""",
                        "images": [img_base64],
                        "stream": False
                    },
                    timeout=300
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("response", "")
                    if not content:
                        content = result.get("message", {}).get("content", "")
                    if content:
                        all_content.append(f"## 第 {page_num} 页\n\n{content}")
        
        if all_content:
            markdown_content = "\n\n---\n\n".join(all_content)
            
            # 自动保存Markdown文件到docs文件夹
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = os.path.splitext(filename)[0]
                md_filename = f"{base_name}_{timestamp}.md"
                md_filepath = os.path.join(DOCS_DIR, md_filename)
                
                with open(md_filepath, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
            except Exception as save_error:
                print(f"保存Markdown文件失败: {str(save_error)}")
            
            return markdown_content, None
        else:
            return "", "未能识别出任何内容"
            
    except Exception as e:
        return "", f"PDF转Markdown失败: {str(e)}"


# ============ Gradio 界面函数 ============

def get_index_list(max_retries=3):
    """获取索引列表，带重试机制"""
    if not PINECONE_API_KEY:
        return [], "未配置PINECONE_API_KEY"
    
    import time
    
    for attempt in range(max_retries):
        try:
            start_time = time.time()
            pc = Pinecone(api_key=PINECONE_API_KEY)
            indexes = [idx["name"] for idx in pc.list_indexes()]
            elapsed = time.time() - start_time
            print(f"获取索引列表成功，耗时: {elapsed:.2f}秒，索引数: {len(indexes)}")
            return indexes, None
        except Exception as e:
            print(f"获取索引列表失败 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(2)  # 等待2秒后重试
            else:
                return [], f"获取索引列表失败: {str(e)}"


def get_namespace_list(index_name: str):
    """获取命名空间列表"""
    if not index_name or index_name == "请先创建索引":
        return ["default"]
    
    try:
        import time
        start_time = time.time()
        
        index, error = get_pinecone_index(index_name)
        if error:
            print(f"获取索引失败: {error}")
            return ["default"]
        
        stats = index.describe_index_stats()
        namespaces = []
        
        # 获取命名空间数据
        ns_data = None
        if stats:
            if hasattr(stats, 'namespaces'):
                ns_data = stats.namespaces
            elif hasattr(stats, '__dict__'):
                ns_data = stats.__dict__.get('namespaces')
            elif isinstance(stats, dict):
                ns_data = stats.get('namespaces')
        
        if ns_data and isinstance(ns_data, dict):
            for ns in ns_data.keys():
                if ns and str(ns).strip():
                    namespaces.append(str(ns))
        
        # 添加default选项
        if "default" not in namespaces:
            namespaces.insert(0, "default")
        
        result = sorted(list(set(namespaces)))
        elapsed = time.time() - start_time
        print(f"获取命名空间列表成功，耗时: {elapsed:.2f}秒，命名空间: {result}")
        return result
    except Exception as e:
        print(f"获取命名空间列表失败: {str(e)}")
        return ["default"]


def refresh_namespaces(index_name: str):
    """刷新命名空间列表"""
    if not index_name or index_name == "请先创建索引":
        return gr.update(choices=["default"], value="default")
    
    namespaces = get_namespace_list(index_name)
    return gr.update(choices=namespaces, value=namespaces[0] if namespaces else "default")


def scan_folder(folder_path: str):
    """递归扫描文件夹，返回所有支持的文件路径列表"""
    supported_extensions = {'.txt', '.md', '.doc', '.docx', '.xlsx', '.pdf', '.jpg', '.jpeg', '.png', '.bmp'}
    files = []
    
    if not folder_path or not os.path.exists(folder_path):
        return []
    
    try:
        for root, dirs, filenames in os.walk(folder_path):
            for filename in filenames:
                ext = os.path.splitext(filename)[1].lower()
                if ext in supported_extensions:
                    full_path = os.path.join(root, filename)
                    files.append(full_path)
        
        # 按路径排序，保持一定顺序
        files.sort()
        return files
    except Exception as e:
        print(f"扫描文件夹失败: {e}")
        return []


def upload_files(files, folder_path, index_name, namespace, max_chars, overlap, ollama_url, ocr_model, 
                 embedding_api_key, embedding_base_url, embedding_model, embedding_dimension):
    """上传文件到Pinecone，支持文件选择和文件夹选择（生成器模式，实时显示进度）"""
    
    # 优先使用文件夹路径（如果填写了）
    if folder_path and folder_path.strip():
        folder_path = folder_path.strip()
        if not os.path.exists(folder_path):
            yield f"❌ 文件夹不存在: {folder_path}"
            return
        
        files = scan_folder(folder_path)
        if not files:
            yield f"❌ 文件夹中没有支持的文件类型\n支持的格式: .txt, .md, .doc, .docx, .xlsx, .pdf, .jpg, .jpeg, .png, .bmp"
            return
    elif files:
        # 使用选择的文件
        pass
    else:
        yield "❌ 请选择文件或输入文件夹路径"
        return
    
    if not index_name:
        yield "❌ 请先选择一个索引"
        return
    
    # 使用自定义命名空间或选择的命名空间
    final_namespace = namespace.strip() if namespace and namespace.strip() else "default"
    
    # 使用自定义嵌入模型配置
    api_key = embedding_api_key.strip() if embedding_api_key and embedding_api_key.strip() else None
    base_url = embedding_base_url.strip() if embedding_base_url and embedding_base_url.strip() else None
    model = embedding_model.strip() if embedding_model and embedding_model.strip() else None
    dimension = int(embedding_dimension) if embedding_dimension else None
    
    index, error = get_pinecone_index(index_name, dimension=dimension)
    if error:
        yield f"❌ {error}"
        return
    
    results = []
    total_docs = 0
    total_vectors = 0
    total_files = len(files)
    
    results.append(f"📁 共发现 {total_files} 个文件")
    results.append(f"{'='*50}")
    yield "\n".join(results)
    
    for file_idx, file_obj in enumerate(files, 1):
        try:
            # 处理文件夹模式（file_obj是字符串路径）和文件模式（file_obj是文件对象）
            if isinstance(file_obj, str):
                # 文件夹模式：file_obj是文件路径字符串
                filename = os.path.basename(file_obj)
                results.append(f"\n📄 [{file_idx}/{total_files}] 正在处理: {filename}")
                yield "\n".join(results)
                
                # 读取文件内容
                with open(file_obj, 'rb') as f:
                    file_content = f.read()
            else:
                # 文件模式：file_obj是Gradio文件对象
                filename = os.path.basename(file_obj.name)
                results.append(f"\n📄 [{file_idx}/{total_files}] 正在处理: {filename}")
                yield "\n".join(results)
                
                # 读取文件内容
                file_content = file_obj.read() if hasattr(file_obj, 'read') else b''
            
            # 判断文件类型
            ext = os.path.splitext(filename)[1].lower()
            
            # PDF文件：先转Markdown，再用Markdown内容上传（使用三层级切分）
            if ext == '.pdf':
                results.append(f"  📝 PDF转Markdown中...")
                yield "\n".join(results)
                
                markdown_content, error = pdf_to_markdown_with_ollama(file_content, ollama_url, ocr_model, filename)
                
                if error:
                    results.append(f"  ❌ PDF转Markdown失败: {error}")
                    yield "\n".join(results)
                    continue
                
                results.append(f"  ✅ Markdown已保存到 docs 文件夹")
                yield "\n".join(results)
                
                # 使用PDF专用三层级切分
                if int(max_chars) != 512:
                    results.append(f"  ⚡ 检测到PDF，已自动启用智能优化参数: 忽略全局设置，强制使用 chunk_size=512, overlap=100 以获得最佳检索效果")
                else:
                    results.append(f"  ⚡ 检测到PDF，已自动启用智能优化参数: chunk_size=512, overlap=100")
                yield "\n".join(results)
                
                pdf_chunks = split_pdf_markdown(markdown_content, chunk_size=512, chunk_overlap=100)
                
                if not pdf_chunks:
                    results.append(f"  ⚠️ PDF内容切分后为空，跳过")
                    yield "\n".join(results)
                    continue
                
                total_docs += 1
                
                # 统计信息
                text_chunks = [c for c in pdf_chunks if c[2] == 'text']
                table_chunks = [c for c in pdf_chunks if c[2] == 'table']
                results.append(f"  ✓ 切分完成: {len(pdf_chunks)} 个块 (文本:{len(text_chunks)}, 表格:{len(table_chunks)})")
                yield "\n".join(results)
                
                # 生成向量并上传
                vectors_to_upsert = []
                for batch_id, batch_chunks in enumerate(batched(pdf_chunks, batch_size=10), start=1):
                    results.append(f"  🔄 生成向量批次 {batch_id}...")
                    yield "\n".join(results)
                    
                    # 提取纯文本内容用于生成向量
                    batch_texts = [chunk[1] for chunk in batch_chunks]
                    
                    embeddings, error = embed_texts(
                        batch_texts, 
                        model=model, 
                        api_key=api_key, 
                        base_url=base_url
                    )
                    if error:
                        results.append(f"  ❌ 嵌入失败: {error}")
                        yield "\n".join(results)
                        continue
                    
                    for i, ((header, content, content_type), embedding) in enumerate(zip(batch_chunks, embeddings), start=1):
                        vec_id = generate_safe_id(filename, batch_id, i)
                        vectors_to_upsert.append({
                            "id": vec_id,
                            "values": embedding,
                            "metadata": {
                                "source": filename,
                                "chunk_index": i,
                                "text": content,
                                "header": header,
                                "content_type": content_type,
                            },
                        })
                
                if vectors_to_upsert:
                    uploaded = smart_upsert(index, vectors_to_upsert, final_namespace, results)
                    total_vectors += uploaded
                    results.append(f"  ✅ 完成! 上传 {uploaded} 个向量")
                    yield "\n".join(results)
            else:
                # 非PDF文件：使用原来的提取方式
                # 创建临时文件对象
                class TempFile:
                    def __init__(self, name, content):
                        self.name = name
                        self._content = content
                    
                    def read(self):
                        return self._content
                
                temp_file = TempFile(filename, file_content)
                content, error = extract_text_from_file(temp_file, ollama_url, ocr_model)
                
                if error:
                    results.append(f"  ❌ {error}")
                    yield "\n".join(results)
                    continue
                
                if not content or not content.strip():
                    results.append(f"  ⚠️ 文件内容为空，跳过")
                    yield "\n".join(results)
                    continue
                
                total_docs += 1
                
                # 分割文本
                chunks = split_text(content, max_chars=max_chars, overlap=overlap)
                results.append(f"  ✓ 分割成 {len(chunks)} 个文本块")
                yield "\n".join(results)
                
                # 生成向量并上传
                vectors_to_upsert = []
                for batch_id, batch_chunks in enumerate(batched(chunks, batch_size=10), start=1):
                    results.append(f"  🔄 生成向量批次 {batch_id}...")
                    yield "\n".join(results)
                    
                    embeddings, error = embed_texts(
                        batch_chunks, 
                        model=model, 
                        api_key=api_key, 
                        base_url=base_url
                    )
                    if error:
                        results.append(f"  ❌ 嵌入失败: {error}")
                        yield "\n".join(results)
                        continue
                    
                    for i, (chunk_text, embedding) in enumerate(zip(batch_chunks, embeddings), start=1):
                        vec_id = generate_safe_id(filename, batch_id, i)
                        vectors_to_upsert.append({
                            "id": vec_id,
                            "values": embedding,
                            "metadata": {
                                "source": filename,
                                "chunk_index": i,
                                "text": chunk_text,
                            },
                        })
                
                if vectors_to_upsert:
                    uploaded = smart_upsert(index, vectors_to_upsert, final_namespace, results)
                    total_vectors += uploaded
                    results.append(f"  ✅ 完成! 上传 {uploaded} 个向量")
                    yield "\n".join(results)
        
        except Exception as e:
            results.append(f"  ❌ 处理失败: {str(e)}")
            yield "\n".join(results)
    
    results.append(f"\n{'='*50}")
    results.append(f"📊 上传完成!")
    results.append(f"  - 总文件数: {total_files}")
    results.append(f"  - 成功处理: {total_docs}")
    results.append(f"  - 创建向量: {total_vectors}")
    results.append(f"  - 命名空间: {final_namespace}")
    
    yield "\n".join(results)


def search_documents(query, index_name, namespace, top_k=5, 
                     embedding_api_key=None, embedding_base_url=None, embedding_model=None):
    """搜索文档"""
    if not query.strip():
        return "请输入搜索内容"
    
    if not index_name:
        return "请先选择一个索引"
    
    final_namespace = namespace.strip() if namespace and namespace.strip() else "default"
    
    # 使用自定义嵌入模型配置
    api_key = embedding_api_key.strip() if embedding_api_key and embedding_api_key.strip() else None
    base_url = embedding_base_url.strip() if embedding_base_url and embedding_base_url.strip() else None
    model = embedding_model.strip() if embedding_model and embedding_model.strip() else None
    
    index, error = get_pinecone_index(index_name)
    if error:
        return error
    
    # 获取查询的向量
    embeddings, error = embed_texts([query], model=model, api_key=api_key, base_url=base_url)
    if error:
        return error
    
    if not embeddings:
        return "无法生成查询向量"
    
    try:
        results = index.query(
            vector=embeddings[0],
            top_k=top_k,
            namespace=final_namespace,
            include_metadata=True
        )
        
        if not results.matches:
            return f"在命名空间 '{final_namespace}' 中未找到相关文档"
        
        output = [f"🔍 搜索结果 (命名空间: {final_namespace})\n"]
        for i, match in enumerate(results.matches, 1):
            source = match.metadata.get("source", "未知")
            text = match.metadata.get("text", "")[:300]
            score = match.score
            output.append(f"{i}. 📄 {source} (相似度: {score:.3f})")
            output.append(f"   {text}...\n")
        
        return "\n".join(output)
    except Exception as e:
        return f"搜索失败: {str(e)}"


def scan_zombie_vectors(index_name, namespace, sample_size, filter_text=""):
    """扫描命名空间中存在的文件来源，支持按文件名关键词筛选"""
    if not index_name:
        return "请先选择一个索引"
    
    index, error = get_pinecone_index(index_name)
    if error:
        return error
    
    try:
        stats = index.describe_index_stats()
        total_vectors = stats['namespaces'].get(namespace, {}).get('vector_count', 0)
        
        if total_vectors == 0:
            return f"📊 命名空间 '{namespace}' 是空的，没有任何向量。"
        
        dimension = EMBEDDING_DIMENSION
        dummy_vector = [0.0] * dimension
        
        query_result = index.query(
            vector=dummy_vector,
            top_k=min(sample_size, 10000),
            namespace=namespace,
            include_metadata=True
        )
        
        found_sources = {}
        for match in query_result.matches:
            if match.metadata and "source" in match.metadata:
                source = match.metadata["source"]
                found_sources[source] = found_sources.get(source, 0) + 1
        
        if not found_sources:
            return f"📊 总向量数: {total_vectors}\n⚠️ 未发现包含 source 字段的向量"
        
        # 处理筛选关键词（按行分割）
        active_filters = []
        if filter_text and filter_text.strip():
            active_filters = [line.strip() for line in filter_text.strip().split('\n') if line.strip()]
        
        # 如果有筛选条件，过滤文件列表
        filtered_sources = {}
        if active_filters:
            for source, count in found_sources.items():
                # 只要匹配任意一个关键词就包含（不区分大小写）
                if any(keyword.lower() in source.lower() for keyword in active_filters):
                    filtered_sources[source] = count
        else:
            filtered_sources = found_sources
        
        result_text = f"📊 命名空间 '{namespace}' 统计:\n"
        result_text += f"总向量数: {total_vectors}\n"
        result_text += f"采样数: {len(query_result.matches)}\n"
        result_text += f"发现文件数: {len(found_sources)}\n"
        
        if active_filters:
            result_text += f"\n🔍 筛选条件（{len(active_filters)}个）:\n"
            for i, kw in enumerate(active_filters, 1):
                result_text += f"  {i}. {kw}\n"
            result_text += f"\n✅ 匹配文件数: {len(filtered_sources)}\n"
        
        result_text += "\n" + "=" * 50 + "\n"
        result_text += "📁 已上传的文件列表:\n"
        result_text += "=" * 50 + "\n"
        
        if not filtered_sources:
            result_text += "⚠️ 没有匹配筛选条件的文件\n"
        else:
            sorted_sources = sorted(filtered_sources.items(), key=lambda x: x[1], reverse=True)
            for source, count in sorted_sources:
                result_text += f"  📄 {source} ({count} 个向量)\n"
        
        result_text += "\n" + "=" * 50 + "\n"
        result_text += "💡 提示：如需删除某些文件，请手动将文件名复制到下方输入框\n"
        result_text += "   （每行一个文件名，仅删除上传失败的文件）\n"
        
        return result_text
        
    except Exception as e:
        return f"扫描失败: {str(e)}"


def delete_file_vectors(index_name, namespace, files_to_delete):
    """删除指定文件的向量"""
    if not index_name:
        return "请先选择一个索引"
    
    if not files_to_delete:
        return "请选择要删除的文件"
    
    index, error = get_pinecone_index(index_name)
    if error:
        return error
    
    if isinstance(files_to_delete, str):
        files_str = files_to_delete.strip()
        if files_str.startswith('[') and files_str.endswith(']'):
            import ast
            try:
                files = ast.literal_eval(files_str)
                if isinstance(files, list):
                    files = [f.strip() for f in files if f.strip()]
                else:
                    files = []
            except:
                files = [f.strip() for f in files_str.split('\n') if f.strip()]
        else:
            files = [f.strip() for f in files_str.split('\n') if f.strip()]
    else:
        files = list(files_to_delete)
    
    if not files:
        return "没有有效的文件名"
    
    results = []
    success_count = 0
    
    for filename in files:
        try:
            index.delete(
                filter={"source": {"$eq": filename}},
                namespace=namespace
            )
            results.append(f"✅ 已删除: {filename}")
            success_count += 1
        except Exception as e:
            results.append(f"❌ 删除失败 {filename}: {str(e)}")
    
    summary = f"\n{'='*50}\n📊 清理完成: 成功删除 {success_count}/{len(files)} 个文件的向量\n"
    
    return "\n".join(results) + summary


def convert_pdf_to_markdown(pdf_file, ollama_url, ocr_model):
    """PDF转Markdown的Gradio界面函数"""
    if pdf_file is None:
        return "请先上传PDF文件"
    
    try:
        # 读取PDF文件
        if hasattr(pdf_file, 'read'):
            pdf_bytes = pdf_file.read()
            filename = pdf_file.name if hasattr(pdf_file, 'name') else "document.pdf"
        else:
            return "无效的文件对象"
        
        # 调用转换函数
        markdown_content, error = pdf_to_markdown_with_ollama(pdf_bytes, ollama_url, ocr_model, filename)
        
        if error:
            return f"转换失败: {error}"
        
        return markdown_content
    except Exception as e:
        return f"转换过程出错: {str(e)}"


# ============ 创建 Gradio 界面 ============

def create_ui():
    # 尝试获取索引列表，失败时使用默认值
    try:
        indexes, error = get_index_list()
    except Exception as e:
        print(f"获取索引列表异常: {e}")
        indexes, error = [], f"Pinecone连接失败: {str(e)}"
    
    # 启动时不获取命名空间，避免网络问题导致启动失败
    initial_namespaces = ["default"]
    
    with gr.Blocks(title="PineDocs - 文档向量管理系统", css="""
        .container { max-width: 1200px; margin: 0 auto; }
        .header { text-align: center; margin-bottom: 2rem; }
        .header h1 { color: #1f77b4; }
        .section { margin: 1rem 0; padding: 1rem; border: 1px solid #e0e0e0; border-radius: 8px; }
    """) as demo:
        
        gr.Markdown("""
        # 🌲 PineDocs
        ## 文档向量管理系统
        """)
        
        if error:
            gr.Warning(error)
        
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("### ⚙️ 配置")
                
                index_dropdown = gr.Dropdown(
                    choices=indexes if indexes else ["请先创建索引"],
                    value=indexes[0] if indexes else "请先创建索引",
                    label="选择索引",
                    interactive=True
                )
                
                refresh_btn = gr.Button("🔄 刷新命名空间", size="sm")
                
                namespace_dropdown = gr.Dropdown(
                    choices=initial_namespaces,
                    value=initial_namespaces[0] if initial_namespaces else "default",
                    label="选择已有命名空间",
                    interactive=True,
                    info="只显示有数据的命名空间"
                )
                
                custom_namespace = gr.Textbox(
                    label="输入命名空间（新建或已有）",
                    placeholder="输入命名空间名称",
                    value="default",
                    info="上传文档时会自动创建，搜索时会自动查找"
                )
                
                ollama_url = gr.Textbox(
                    label="Ollama URL",
                    value="http://localhost:11434",
                    info="用于OCR识别的本地服务地址"
                )
                
                ocr_model = gr.Dropdown(
                    choices=["my-glm-ocr:latest", "my-PaddleOCR-VL:0.9b"],
                    value="my-glm-ocr:latest",
                    label="OCR模型"
                )
                
                gr.Markdown("---")
                gr.Markdown("**嵌入模型设置**")
                
                embedding_api_key = gr.Textbox(
                    label="嵌入模型 API Key",
                    value=EMBEDDING_API_KEY if EMBEDDING_API_KEY else "",
                    placeholder="输入API密钥",
                    info="支持 OpenAI 兼容格式",
                    type="password"
                )
                
                embedding_base_url = gr.Textbox(
                    label="嵌入模型 Base URL",
                    value=EMBEDDING_BASE_URL,
                    placeholder="https://api.example.com/v1",
                    info="OpenAI 兼容的 API 基础URL"
                )
                
                embedding_model = gr.Textbox(
                    label="嵌入模型名称",
                    value=EMBEDDING_MODEL,
                    placeholder="text-embedding-3-small",
                    info="模型名称，如 text-embedding-3-small"
                )
                
                embedding_dimension = gr.Number(
                    label="向量维度",
                    value=EMBEDDING_DIMENSION,
                    precision=0,
                    info="模型的输出维度（如 1536, 3072）"
                )
            
            with gr.Column(scale=2):
                with gr.Tab("📤 上传文档"):
                    gr.Markdown("### 上传文件到向量数据库")
                    
                    gr.Markdown("**方式一：选择文件（支持多选）**")
                    file_input = gr.File(
                        label="选择文件",
                        file_count="multiple",
                        file_types=[".txt", ".md", ".doc", ".docx", ".xlsx", ".pdf", ".jpg", ".jpeg", ".png", ".bmp"]
                    )
                    
                    gr.Markdown("**方式二：输入文件夹路径（自动扫描子文件夹）**")
                    folder_input = gr.Textbox(
                        label="文件夹路径",
                        placeholder="输入文件夹完整路径，如: D:\\Documents\\MyFiles，留空则不使用此方式"
                    )
                    
                    with gr.Row():
                        max_chars = gr.Slider(
                            minimum=100,
                            maximum=5000,
                            value=900,
                            step=100,
                            label="每块最大字符数"
                        )
                        overlap = gr.Slider(
                            minimum=0,
                            maximum=1000,
                            value=100,
                            step=50,
                            label="重叠字符数"
                        )
                    
                    upload_btn = gr.Button("🚀 开始上传", variant="primary", size="lg")
                    upload_output = gr.Textbox(
                        label="上传结果",
                        lines=15,
                        max_lines=30,
                        interactive=False
                    )
                
                with gr.Tab("🔍 搜索文档"):
                    gr.Markdown("### 语义搜索")
                    
                    query_input = gr.Textbox(
                        label="搜索内容",
                        placeholder="输入您想搜索的内容...",
                        lines=3
                    )
                    
                    top_k = gr.Slider(
                        minimum=1,
                        maximum=20,
                        value=5,
                        step=1,
                        label="返回结果数量"
                    )
                    
                    search_btn = gr.Button("🔍 搜索", variant="primary")
                    search_output = gr.Textbox(
                        label="搜索结果",
                        lines=20,
                        max_lines=40,
                        interactive=False
                    )
                
                with gr.Tab("🗑️ 管理"):
                    gr.Markdown("### 🧹 残留数据清理")
                    gr.Markdown("""
                    **功能说明:**
                    - **扫描**: 查看当前命名空间中已上传的文件列表
                    - **清理**: 删除指定文件的向量（用于处理上传中断的残留数据）
                    
                    ⚠️ **注意**: 如果上传中断，前面的批次已经存入数据库，需要手动清理后再重传！
                    """)
                    
                    with gr.Row():
                        scan_namespace_dropdown = gr.Dropdown(
                            choices=initial_namespaces,
                            value=initial_namespaces[0] if initial_namespaces else "default",
                            label="命名空间",
                            interactive=True,
                            allow_custom_value=True
                        )
                        scan_sample_size = gr.Slider(
                            minimum=100,
                            maximum=10000,
                            value=2000,
                            step=100,
                            label="采样数量",
                            info="采样越多越准确，但速度越慢"
                        )
                    
                    with gr.Accordion("📁 按文件名筛选（可选）", open=False):
                        gr.Markdown("输入文件名关键词进行筛选，支持多条（每行一个关键词）")
                        
                        filename_filter_input = gr.Textbox(
                            label="文件名包含（每行一个）",
                            placeholder="输入文件名关键词，每行一个\n例如:\n丽水市\n绿色建筑\n规划",
                            lines=4,
                            interactive=True
                        )
                    
                    scan_btn = gr.Button("🔍 扫描文件列表", variant="primary")
                    scan_output = gr.Textbox(
                        label="扫描结果",
                        lines=15,
                        max_lines=30,
                        interactive=False
                    )
                    
                    files_to_delete = gr.Textbox(
                        label="要删除的文件（每行一个文件名）",
                        placeholder="从扫描结果中复制需要删除的文件名，每行一个\n例如:\n文件1.pdf\n文件2.docx",
                        lines=5,
                        interactive=True
                    )
                    
                    delete_btn = gr.Button("🗑️ 删除选中文件", variant="stop")
                    delete_output = gr.Textbox(label="删除结果", interactive=False)
                
                with gr.Tab("📄 PDF转Markdown"):
                    gr.Markdown("### PDF转Markdown（使用OCR）")
                    
                    pdf_input = gr.File(
                        label="选择PDF文件",
                        file_types=[".pdf"]
                    )
                    
                    gr.Markdown("""
                    **使用说明:**
                    1. 确保Ollama已安装并运行
                    2. 确保已安装所选模型: `ollama pull {model}`
                    3. 上传PDF文件
                    4. 点击转换按钮
                    5. 转换后的Markdown会自动保存到docs文件夹
                    
                    **模型说明:**
                    - **my-PaddleOCR-VL:0.9b**: PaddleOCR-VL-1.5，文档识别效果更好的模型
                    - **my-glm-ocr:latest**: GLM-OCR，通用OCR模型
                    """.format(model=ocr_model.value if hasattr(ocr_model, 'value') else "my-glm-ocr:latest"))
                    
                    convert_btn = gr.Button("🔄 开始转换", variant="primary", size="lg")
                    
                    markdown_output = gr.Textbox(
                        label="转换结果（Markdown格式）",
                        lines=25,
                        max_lines=50,
                        interactive=False
                    )
        
        # 事件绑定
        refresh_btn.click(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[namespace_dropdown]
        )
        
        refresh_btn.click(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[scan_namespace_dropdown]
        )
        
        index_dropdown.change(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[namespace_dropdown]
        )
        
        index_dropdown.change(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[scan_namespace_dropdown]
        )
        
        upload_btn.click(
            fn=upload_files,
            inputs=[file_input, folder_input, index_dropdown, custom_namespace, max_chars, overlap, ollama_url, ocr_model,
                   embedding_api_key, embedding_base_url, embedding_model, embedding_dimension],
            outputs=[upload_output],
            show_progress=True
        )
        
        search_btn.click(
            fn=search_documents,
            inputs=[query_input, index_dropdown, custom_namespace, top_k,
                   embedding_api_key, embedding_base_url, embedding_model],
            outputs=[search_output]
        )
        
        refresh_btn.click(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[scan_namespace_dropdown]
        )
        
        index_dropdown.change(
            fn=refresh_namespaces,
            inputs=[index_dropdown],
            outputs=[scan_namespace_dropdown]
        )
        
        scan_btn.click(
            fn=scan_zombie_vectors,
            inputs=[index_dropdown, scan_namespace_dropdown, scan_sample_size, filename_filter_input],
            outputs=[scan_output]
        )
        
        delete_btn.click(
            fn=delete_file_vectors,
            inputs=[index_dropdown, scan_namespace_dropdown, files_to_delete],
            outputs=[delete_output]
        )
        
        convert_btn.click(
            fn=convert_pdf_to_markdown,
            inputs=[pdf_input, ollama_url, ocr_model],
            outputs=[markdown_output]
        )
    
    return demo


if __name__ == "__main__":
    demo = create_ui()
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)
